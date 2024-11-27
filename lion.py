import docx 
from collections import Counter
import re #(регулярные выражения) модуль для работы с текстом 
import pandas as pd

doc = docx.Document('lion.docx')

num_par = len(doc.paragraphs) # количество параграфов
word_list = [] #список для слов без повторений 

all_letters = []
rus_let = re.compile(r'[а-яё]', re.IGNORECASE) #регулярное выражение для поиска русских букв

rus_text = [] #список для всего текста (только русские слова)
rus_word = re.compile(r'\b[а-яё]+\b',re.IGNORECASE) #регулярное выражение для поиска русских слов 
prepositions = {'и','в','не','что','он','на','с','как','его','то','к','стр',
                'я','г','но','она','это','из','а','за','так','по','изд',',был',
                'ему','всё','от','о','же','бы','ее','был','у','для','еще','р','вы','они','все','да'} 

for paragraph in doc.paragraphs: #проходимся по всем параграфам в документе
    text = paragraph.text.lower() # приводим текст к нижнему регистру
    words = rus_word.findall(text) # находим все русские слова
    letters = rus_let.findall(text) # находим все русские буквы
    all_letters.extend(letters) #добовляем буквы в общий список

    for word in words:
        if word not in word_list: # проверка на на наличие слова в списке 
            word_list.append(word) #добавления слова 
        if word not in prepositions:
            rus_text.append(word)

word_counts = Counter(rus_text) #подсчитываем частоту встречаемости слов
for word, count in word_counts.items():
    data_word = {f'{word}: {count}'}

letters_counts = Counter(all_letters) #подсчитываем частоту встречаемости слов
df = pd.DataFrame(letters_counts.items(),columns=['Letter','Frequency'])
print(df)

#Закидываю результат в новый doc файл
doc_letter = docx.Document() #создаю новый документ

doc_letter.add_heading('Встречаемость букв в тексте', level=1)#добовляем заголовок

#Добавление данных DataFrame в документ как текст
for index, row in df.iterrows():
    letter_text = [f"{col}: {row[col]}" for col in df.columns]
    doc_letter.add_paragraph(letter_text)

doc_letter.save('встречаемость_букв.docx') # Сохраняем документ

print(len(word_list))
print(num_par)
